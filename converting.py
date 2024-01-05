import subprocess
from pydub import AudioSegment
from PIL import Image as ImagePIL
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import fitz  # PyMuPDF


class Video:

    #video: mp4, mov, avi, wmv

    def convert(input_file, output_file):
        try:
            cmd = ['ffmpeg', '-i', input_file, output_file]
            subprocess.run(cmd, check=True)
            print(f"Conversion successful: {input_file} -> {output_file}")
        except subprocess.CalledProcessError as e:
            print(f"Error converting {input_file} to {output_file}: {e}")
        except FileNotFoundError:
            print("ffmpeg command not found. Make sure ffmpeg is installed and in your system's PATH.")
         
# Video.convert('input.mov', 'output.mp4')
        


class Music:

    # music: mp3, wav, aac

    def convert(input_file, output_file):
        x = input_file.split('.')[-1]
        y = output_file.split('.')[-1]   
        audio = AudioSegment.from_file(input_file, format=x)
        audio.export(output_file, format=y)

# Music.convert('September.mp3', 'September.wav')



class Text:
    
    # text: pdf, docx, txt

    def convert(input_file, output_file):
        extension = input_file.split('.')[-1].lower()
        extension_output = output_file.split('.')[-1].lower()

        if extension == 'pdf':

            if extension_output == 'txt':
                try:
                    doc = fitz.open(input_file)
                    text = ''
                    for page_num in range(doc.page_count):
                        page = doc[page_num]
                        text += page.get_text()

                    with open(output_file, 'w', encoding='utf-8') as text_file:
                        text_file.write(text)

                    print(f"Conversion successful: {input_file} -> {output_file}")

                except Exception as e:
                    print(f"Error converting {input_file} to {output_file}: {e}")


            if extension_output == 'docx':
                pdf_document = fitz.open(input_file)
                text = ''

                for page in pdf_document:
                    text += page.get_text()

                pdf_document.close()
                with open(output_file, 'w', encoding='utf-8') as file:
                    file.write(text)


        elif extension == 'docx':

            if extension_output == 'txt':
                try:
                    doc = Document(input_file)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    with open(output_file, 'w', encoding='utf-8') as text_file:
                        text_file.write(text)

                    print(f"Conversion successful: {input_file} -> {output_file}")

                except Exception as e:
                    print(f"Error converting {input_file} to {output_file}: {e}")

            
            if extension_output == 'pdf':
                try: 
                    doc = Document(input_file)
                    pdf = canvas.Canvas(output_file, pagesize=letter)
                    width, height = letter

                    y = height - 50
                    for para in doc.paragraphs:
                        pdf.drawString(50, y, para.text)
                        y -= 20

                    pdf.save()

                except Exception as e:
                    print(f"Error converting {input_file} to {output_file}: {e}")


        elif extension == 'txt':

            if extension_output == 'pdf':
                try:
                    # Convert text to PDF using reportlab
                    pdf_canvas = canvas.Canvas(output_file)
                    with open(input_file, 'r', encoding='utf-8') as text_file:
                        lines = text_file.readlines()
                        y_coordinate = 800  # Starting y-coordinate for the first line

                        for line in lines:
                            pdf_canvas.drawString(100, y_coordinate, line.strip())
                            y_coordinate -= 12  # Adjust this value for spacing

                    pdf_canvas.save()
                    print(f"Conversion successful: {input_file} -> {output_file}")

                except Exception as e:
                    print(f"Error converting {input_file} to {output_file}: {e}")
        
            if extension_output == 'docx':
                try:
                    # Convert text to DOCX using python-docx
                    doc = Document()
                    with open(input_file, 'r', encoding='utf-8') as text_file:
                        lines = text_file.readlines()
                        for line in lines:
                            doc.add_paragraph(line.strip())
                    doc.save(output_file)
                    print(f"Conversion successful: {input_file} -> {output_file}")
                except Exception as e:
                    print(f"Error converting {input_file} to {output_file}: {e}")
            
# Text.convert('SampleProba.docx', 'probaa.txt')



class Image:

    # images: jpg, png, gif, svg, bmp, tiff, webp, ico

    def convert(input_file, output_file):
        im = ImagePIL.open(input_file)
        rgb_im = im.convert('RGB')
        rgb_im.save(output_file)

# Image.convert('stevie-steve.webp', 'stevie-steve.jpg')



